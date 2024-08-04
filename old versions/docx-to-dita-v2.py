import tkinter as tk
from tkinter import filedialog, messagebox, Menu
import json
import xml.etree.ElementTree as ET
from docx import Document
import xml.dom.minidom
import os

# Global variable for storing keyword replacements
keyword_replacements = {}

def save_preferences():
    global keyword_replacements
    with open('preferences.json', 'w') as f:
        json.dump(keyword_replacements, f)
    messagebox.showinfo("Preferences Saved", "Keyword replacements saved successfully.")

def open_preferences_dialog():
    global keyword_replacements
    
    # Preferences dialog setup
    preferences_window = tk.Toplevel(root)
    preferences_window.title("Keyword Replacements")
    
    # Label and Entry for keyword replacements
    tk.Label(preferences_window, text="Keyword : XML Tag ID").grid(row=0, column=0, padx=10, pady=5)
    
    keyword_entry = tk.Entry(preferences_window, width=50)
    keyword_entry.grid(row=1, column=0, padx=10, pady=5)
    
    # Save button
    def add_keyword():
        keyword_entry_text = keyword_entry.get().strip()
        if keyword_entry_text:
            # Split the entry into keyword and XML tag ID
            try:
                keyword, xml_tag_id = [part.strip() for part in keyword_entry_text.split(':')]
                keyword_replacements[keyword] = xml_tag_id
                keyword_entry.delete(0, tk.END)
                messagebox.showinfo("Keyword Added", f"Keyword '{keyword}' added successfully.")
            except ValueError:
                messagebox.showerror("Error", "Please use the format 'Keyword : XML Tag ID'.")
        else:
            messagebox.showerror("Error", "Please enter a keyword and XML Tag ID.")
    
    tk.Button(preferences_window, text="Add Keyword", command=add_keyword).grid(row=2, column=0, padx=10, pady=10)
    
    # Load existing preferences
    for keyword, xml_tag_id in keyword_replacements.items():
        tk.Label(preferences_window, text=f"{keyword} : {xml_tag_id}").grid(sticky="w", padx=10, pady=5)

    # Save preferences button
    tk.Button(preferences_window, text="Save Preferences", command=save_preferences).grid(row=3, column=0, padx=10, pady=10)

def docx_to_dita_task(docx_path, dita_path, task_id):
    doc = Document(docx_path)
    
    # Create the root element of the DITA task document with task_id attribute
    root = ET.Element('task', id=task_id)
    
    # Add the title to the DITA document
    title = ET.SubElement(root, 'title')
    title.text = doc.paragraphs[0].text  # Assuming the first paragraph is the title
    
    # Add task body content
    task_body = ET.SubElement(root, 'taskbody')
    
    # Process paragraphs to identify steps and substeps
    steps = ET.SubElement(task_body, 'steps')
    
    current_step = None
    current_substeps = None
    
    # Iterate through paragraphs
    for para in doc.paragraphs:
        if para.style.name == 'List Paragraph':
            # New main step
            current_step = ET.SubElement(steps, 'step')
            step_cmd = ET.SubElement(current_step, 'cmd')
            step_cmd.text = para.text.strip()
        elif para.style.name == 'List Number 2':
            # New substep
            if current_step is not None:
                if current_substeps is None:
                    current_substeps = ET.SubElement(current_step, 'substeps')
                substep = ET.SubElement(current_substeps, 'substep')
                substep_cmd = ET.SubElement(substep, 'cmd')
                substep_cmd.text = para.text.strip()
        else:
            # Regular paragraphs are added as step info
            if current_step is not None:
                step_info = ET.SubElement(current_step, 'info')
                
                # Replace keywords with <keyword id="USER_INPUT"></keyword> content
                para_text = para.text.strip()
                for keyword, xml_tag_id in keyword_replacements.items():
                    print(keyword_replacements.items())
                    para_text = para_text.replace(keyword, f'<keyword keyref="{xml_tag_id}"></keyword>')
                
                step_info.text = para_text
    
    # Process images (assuming this part remains unchanged)
    for image in doc.inline_shapes:
        if image.type == 3:  # Inline shapes of type 3 are images
            if include_images.get():
                if ask_for_image_paths.get():
                    # Ask for image path and name
                    image_path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG files", "*.png"), ("JPEG files", "*.jpg"), ("All files", "*.*")])
                    if not image_path:
                        continue
                else:
                    # Default image path and name
                    image_path = f"images/{image._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name}.png"
                
                # Save image to disk
                image._inline.graphic.graphicData.pic.save(image_path)
                
                # Create image tag with optional title attribute
                img_tag = ET.SubElement(current_step, 'image', href=image_path)
                img_tag.set('title', '')
    
    # Create a new XML tree
    tree = ET.ElementTree(root)
    
    # Convert XML to string without XML declaration
    xml_str = ET.tostring(root, encoding='utf-8', method='xml')
    
    # Parse XML string for pretty printing without adding XML declaration
    dom = xml.dom.minidom.parseString(xml_str)
    pretty_xml_str = dom.toprettyxml(indent='  ')
    
    # Remove the first line containing the unwanted XML declaration
    pretty_xml_str = '\n'.join(pretty_xml_str.split('\n')[1:])
    
    # Insert DOCTYPE declaration and XML declaration manually
    xml_declaration = '<?xml version="1.0" encoding="UTF-8"?>\n'
    doctype = '<!DOCTYPE task PUBLIC "-//OASIS//DTD DITA Task//EN" "task.dtd">\n'
    final_xml_str = xml_declaration + doctype + pretty_xml_str
    
    # Write formatted XML to the .dita file
    with open(dita_path, 'w', encoding='utf-8') as f:
        f.write(final_xml_str)


def browse_file():
    filename = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if filename:
        input_path_entry.delete(0, tk.END)
        input_path_entry.insert(0, filename)

def convert_file():
    input_path = input_path_entry.get()
    output_path = output_path_entry.get()
    task_id = task_id_entry.get()
    
    if not input_path:
        messagebox.showerror("Error", "Please select an input file.")
        return
    
    if not output_path:
        messagebox.showerror("Error", "Please specify an output file.")
        return
    
    if not task_id:
        messagebox.showerror("Error", "Please enter a task ID.")
        return
    
    if not input_path.lower().endswith('.docx'):
        messagebox.showerror("Error", "Input file must be a .docx file.")
        return
    
    if not output_path.lower().endswith('.dita'):
        output_path += '.dita'
    
    try:
        docx_to_dita_task(input_path, output_path, task_id)
        messagebox.showinfo("Success", f"Conversion completed successfully. Output saved to {output_path}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred during conversion:\n{str(e)}")

# GUI setup
root = tk.Tk()
root.title("DOCX to DITA Converter")

# Menu bar
menu_bar = Menu(root)
root.config(menu=menu_bar)

# File menu
file_menu = Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="File", menu=file_menu)
file_menu.add_command(label="Open Preferences", command=open_preferences_dialog)
file_menu.add_separator()
file_menu.add_command(label="Exit", command=root.quit)

# Input File
input_path_label = tk.Label(root, text="Input .docx file:")
input_path_label.grid(row=0, column=0, padx=10, pady=5, sticky=tk.W)

input_path_entry = tk.Entry(root, width=50)
input_path_entry.grid(row=0, column=1, padx=10, pady=5)

browse_button = tk.Button(root, text="Browse", command=browse_file)
browse_button.grid(row=0, column=2, padx=10, pady=5)

# Output File
output_path_label = tk.Label(root, text="Output .dita file:")
output_path_label.grid(row=1, column=0, padx=10, pady=5, sticky=tk.W)

output_path_entry = tk.Entry(root, width=50)
output_path_entry.grid(row=1, column=1, padx=10, pady=5)

# Task ID
task_id_label = tk.Label(root, text="Task ID:")
task_id_label.grid(row=2, column=0, padx=10, pady=5, sticky=tk.W)

task_id_entry = tk.Entry(root, width=50)
task_id_entry.grid(row=2, column=1, padx=10, pady=5)

# Checkbox for including images
include_images = tk.BooleanVar()
include_images_checkbutton = tk.Checkbutton(root, text="Include Images", variable=include_images)
include_images_checkbutton.grid(row=3, column=0, columnspan=2, padx=10, pady=5, sticky=tk.W)

# Checkbox for asking for each image path
ask_for_image_paths = tk.BooleanVar()
ask_for_image_paths_checkbutton = tk.Checkbutton(root, text="Ask for Each Image Path", variable=ask_for_image_paths)
ask_for_image_paths_checkbutton.grid(row=4, column=0, columnspan=2, padx=10, pady=5, sticky=tk.W)

# Convert Button
convert_button = tk.Button(root, text="Convert", command=convert_file)
convert_button.grid(row=5, column=1, padx=10, pady=10)

# Preferences Button
preferences_button = tk.Button(root, text="Preferences", command=open_preferences_dialog)
preferences_button.grid(row=5, column=2, padx=10, pady=10)

# Load preferences if file exists
preferences_file = 'preferences.json'
if os.path.exists(preferences_file):
    with open(preferences_file, 'r') as f:
        keyword_replacements = json.load(f)

# Start GUI main loop
root.mainloop()
