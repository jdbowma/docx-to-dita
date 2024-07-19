import tkinter as tk
from tkinter import filedialog, messagebox, Menu
import json
import xml.etree.ElementTree as ET
from docx import Document
import xml.dom.minidom
import os

# Global variable for storing keyword replacements
keyword_replacements = {}

def save_preferences():
    global keyword_replacements
    try:
        prefs = preferences_text.get("1.0", tk.END).strip().split('\n')
        keyword_replacements = {}
        for pref in prefs:
            if ':' in pref:
                original, new = map(str.strip, pref.split(':', 1))
                keyword_replacements[original] = new
        
        with open('preferences.json', 'w') as f:
            json.dump(keyword_replacements, f)
        messagebox.showinfo("Preferences Saved", "Keyword replacements saved successfully.")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred while saving preferences:\n{str(e)}")

def open_preferences_dialog():
    global preferences_text
    preferences_window = tk.Toplevel(root)
    preferences_window.title("Keyword Replacements")
    
    tk.Label(preferences_window, text="Specify replacements in the format 'ORIGINALPHRASE : NEWPHRASE'").pack(padx=10, pady=5)
    
    preferences_text = tk.Text(preferences_window, width=80, height=20)
    preferences_text.pack(padx=10, pady=5)
    
    current_prefs = '\n'.join([f"{key} : {value}" for key, value in keyword_replacements.items()])
    preferences_text.insert(tk.END, current_prefs)
    
    tk.Button(preferences_window, text="Save Preferences", command=save_preferences).pack(padx=10, pady=10)

def docx_to_dita_task(docx_path, dita_path, task_id):
    doc = Document(docx_path)
    
    root = ET.Element('task', id=task_id)
    
    title = ET.SubElement(root, 'title')
    title.text = doc.paragraphs[0].text
    
    task_body = ET.SubElement(root, 'taskbody')
    steps = ET.SubElement(task_body, 'steps')
    
    current_step = None
    current_substeps = None
    
    for para in doc.paragraphs[1:]:
        para_text = para.text.strip()
        if check_for_notes.get() and para_text.startswith("Note:"):
            note_content = para_text[5:].strip()
            if prompt_for_notes.get():
                response = messagebox.askyesno("Note Detected", f"Is this a note?\n\n{note_content}")
                if response:
                    note_tag = ET.Element('note')
                    note_tag.text = note_content
                    info_tag = ET.SubElement(current_step if current_step else steps, 'info')
                    info_tag.append(note_tag)
                else:
                    if current_step is not None:
                        step_info = ET.SubElement(current_step, 'info')
                        step_info.text = note_content
                    else:
                        para_tag = ET.SubElement(steps, 'p')
                        para_tag.text = note_content
            else:
                note_tag = ET.Element('note')
                note_tag.text = note_content
                info_tag = ET.SubElement(current_step if current_step else steps, 'info')
                info_tag.append(note_tag)
        elif para.style.name == 'List Paragraph':
            current_step = ET.SubElement(steps, 'step')
            step_cmd = ET.SubElement(current_step, 'cmd')
            step_cmd.text = para_text
        elif para.style.name == 'List Number 2':
            if current_step is not None:
                if current_substeps is None:
                    current_substeps = ET.SubElement(current_step, 'substeps')
                substep = ET.SubElement(current_substeps, 'substep')
                substep_cmd = ET.SubElement(substep, 'cmd')
                substep_cmd.text = para_text
        else:
            if current_step is not None:
                step_info = ET.SubElement(current_step, 'info')
                step_info.text = para_text
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            if include_images.get():
                img = rel.target_part.blob
                image_path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG files", "*.png"), ("JPEG files", "*.jpg"), ("All files", "*.*")])
                if image_path:
                    with open(image_path, 'wb') as f:
                        f.write(img)
                    if current_step is not None:
                        info_tag = ET.SubElement(current_step, 'info')
                        fig_tag = ET.SubElement(info_tag, 'fig')
                        img_tag = ET.SubElement(fig_tag, 'image', href=image_path)
    
    xml_str = ET.tostring(root, encoding='utf-8', method='xml').decode('utf-8')
    for original, new in keyword_replacements.items():
        xml_str = xml_str.replace(original, new)
        print(f"Replaced {original} with {new}")
    root = ET.fromstring(xml_str)
    
    tree = ET.ElementTree(root)
    xml_str = ET.tostring(root, encoding='utf-8', method='xml')
    dom = xml.dom.minidom.parseString(xml_str)
    pretty_xml_str = dom.toprettyxml(indent='  ')
    pretty_xml_str = '\n'.join(pretty_xml_str.split('\n')[1:])
    xml_declaration = '<?xml version="1.0" encoding="UTF-8"?>\n'
    doctype = '<!DOCTYPE task PUBLIC "-//OASIS//DTD DITA Task//EN" "task.dtd">\n'
    final_xml_str = xml_declaration + doctype + pretty_xml_str
    
    with open(dita_path, 'w', encoding='utf-8') as f:
        f.write(final_xml_str)

def browse_file():
    filename = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if filename:
        input_path_entry.delete(0, tk.END)
        input_path_entry.insert(0, filename)

def convert_file():
    input_path = input_path_entry.get()
    output_path = output_path_entry.get()
    task_id = task_id_entry.get()
    
    if not input_path:
        messagebox.showerror("Error", "Please select an input file.")
        return
    
    if not output_path:
        messagebox.showerror("Error", "Please specify an output file.")
        return
    
    if not task_id:
        messagebox.showerror("Error", "Please enter a task ID.")
        return
    
    if not input_path.lower().endswith('.docx'):
        messagebox.showerror("Error", "Input file must be a .docx file.")
        return
    
    if not output_path.lower().endswith('.dita'):
        output_path += '.dita'
    
    try:
        docx_to_dita_task(input_path, output_path, task_id)
        messagebox.showinfo("Success", f"Conversion completed successfully. Output saved to {output_path}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred during conversion:\n{str(e)}")

root = tk.Tk()
root.title("DOCX to DITA Converter")

menu_bar = Menu(root)
root.config(menu=menu_bar)

file_menu = Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="File", menu=file_menu)
file_menu.add_command(label="Open Preferences", command=open_preferences_dialog)
file_menu.add_separator()
file_menu.add_command(label="Exit", command=root.quit)

input_path_label = tk.Label(root, text="Input .docx file:")
input_path_label.grid(row=0, column=0, padx=10, pady=5, sticky=tk.W)

input_path_entry = tk.Entry(root, width=50)
input_path_entry.grid(row=0, column=1, padx=10, pady=5)

browse_button = tk.Button(root, text="Browse", command=browse_file)
browse_button.grid(row=0, column=2, padx=10, pady=5)

output_path_label = tk.Label(root, text="Output .dita file:")
output_path_label.grid(row=1, column=0, padx=10, pady=5, sticky=tk.W)

output_path_entry = tk.Entry(root, width=50)
output_path_entry.grid(row=1, column=1, padx=10, pady=5)

task_id_label = tk.Label(root, text="Task ID:")
task_id_label.grid(row=2, column=0, padx=10, pady=5, sticky=tk.W)

task_id_entry = tk.Entry(root, width=50)
task_id_entry.grid(row=2, column=1, padx=10, pady=5)

include_images = tk.BooleanVar()
include_images_checkbutton = tk.Checkbutton(root, text="Include Images", variable=include_images)
include_images_checkbutton.grid(row=3, column=0, columnspan=2, padx=10, pady=5, sticky=tk.W)

ask_for_image_paths = tk.BooleanVar()
ask_for_image_paths_checkbutton = tk.Checkbutton(root, text="Ask for Each Image Path", variable=ask_for_image_paths)
ask_for_image_paths_checkbutton.grid(row=4, column=0, columnspan=2, padx=10, pady=5, sticky=tk.W)

check_for_notes = tk.BooleanVar()
check_for_notes_checkbutton = tk.Checkbutton(root, text="Check for notes", variable=check_for_notes)
check_for_notes_checkbutton.grid(row=5, column=0, columnspan=2, padx=10, pady=5, sticky=tk.W)

prompt_for_notes = tk.BooleanVar()
prompt_for_notes_checkbutton = tk.Checkbutton(root, text="Prompt for note verification", variable=prompt_for_notes)
prompt_for_notes_checkbutton.grid(row=6, column=0, columnspan=2, padx=10, pady=5, sticky=tk.W)

convert_button = tk.Button(root, text="Convert", command=convert_file)
convert_button.grid(row=7, column=1, padx=10, pady=10)

preferences_button = tk.Button(root, text="Preferences", command=open_preferences_dialog)
preferences_button.grid(row=7, column=2, padx=10, pady=10)

preferences_file = 'preferences.json'
if os.path.exists(preferences_file):
    with open(preferences_file, 'r') as f:
        keyword_replacements = json.load(f)

root.mainloop()
