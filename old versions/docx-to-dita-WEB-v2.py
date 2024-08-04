from flask import Flask, request, jsonify
import os
import json
import xml.etree.ElementTree as ET
from docx import Document
import xml.dom.minidom

app = Flask(__name__)

@app.route('/convert', methods=['POST'])
def convert():
    try:
        task_id = request.form['taskId']
        preferences = request.form['preferences']
        include_images = request.form.get('includeImages') == 'true'
        ask_for_image_paths = request.form.get('askForImagePaths') == 'true'

        keyword_replacements = {}
        for pref in preferences.split('\n'):
            if ':' in pref:
                original, new = map(str.strip, pref.split(':', 1))
                keyword_replacements[original] = new

        input_file = request.files['inputFile']
        input_file.save('input.docx')

        doc = Document('input.docx')
        root = ET.Element('task', id=task_id)
        title = ET.SubElement(root, 'title')
        title.text = doc.paragraphs[0].text
        task_body = ET.SubElement(root, 'taskbody')
        steps = ET.SubElement(task_body, 'steps')
        current_step = None
        current_substeps = None

        for para in doc.paragraphs[1:]:
            if para.style.name == 'List Paragraph':
                current_step = ET.SubElement(steps, 'step')
                step_cmd = ET.SubElement(current_step, 'cmd')
                step_cmd.text = para.text.strip()
            elif para.style.name == 'List Number 2':
                if current_step is not None:
                    if current_substeps is None:
                        current_substeps = ET.SubElement(current_step, 'substeps')
                    substep = ET.SubElement(current_substeps, 'substep')
                    substep_cmd = ET.SubElement(substep, 'cmd')
                    substep_cmd.text = para.text.strip()
            else:
                if current_step is not None:
                    step_info = ET.SubElement(current_step, 'info')
                    step_info.text = para.text.strip()

        xml_str = ET.tostring(root, encoding='utf-8', method='xml').decode('utf-8')
        for original, new in keyword_replacements.items():
            xml_str = xml_str.replace(original, new)
        root = ET.fromstring(xml_str)

        tree = ET.ElementTree(root)
        xml_str = ET.tostring(root, encoding='utf-8', method='xml')
        dom = xml.dom.minidom.parseString(xml_str)
        pretty_xml_str = dom.toprettyxml(indent='  ')
        pretty_xml_str = '\n'.join(pretty_xml_str.split('\n')[1:])
        xml_declaration = '<?xml version="1.0" encoding="UTF-8"?>\n'
        doctype = '<!DOCTYPE task PUBLIC "-//OASIS//DTD DITA Task//EN" "task.dtd">\n'
        final_xml_str = xml_declaration + doctype + pretty_xml_str

        with open('output.dita', 'w', encoding='utf-8') as f:
            f.write(final_xml_str)

        return jsonify(success=True, message="Conversion completed successfully.")

    except Exception as e:
        return jsonify(success=False, message=str(e))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
