from docx import Document
import xml.etree.ElementTree as ET
import xml.dom.minidom
import os
import sys

def docx_to_dita_task(docx_path, dita_path, task_id):
    doc = Document(docx_path)
    
    # Create the root element of the DITA task document with task_id attribute
    root = ET.Element('task', id=task_id)
    
    # Add the title to the DITA document
    title = ET.SubElement(root, 'title')
    title.text = doc.paragraphs[0].text  # Assuming the first paragraph is the title
    
    # Add task body content
    task_body = ET.SubElement(root, 'taskbody')
    
    # Process paragraphs to identify steps
    steps = ET.SubElement(task_body, 'steps')
    
    current_step = None
    
    for para in doc.paragraphs[1:]:  # Skipping the first paragraph (title)
        if para.style.name == 'List Paragraph':
            # New main step
            current_step = ET.SubElement(steps, 'step')
            step_cmd = ET.SubElement(current_step, 'cmd')
            step_cmd.text = para.text.strip()
        else:
            # Regular paragraphs are added as step info
            if current_step is not None:
                step_info = ET.SubElement(current_step, 'info')
                step_info.text = para.text.strip()
    
    # Create a new XML tree
    tree = ET.ElementTree(root)
    
    # Convert XML to string without XML declaration
    xml_str = ET.tostring(root, encoding='utf-8', method='xml')
    
    # Parse XML string for pretty printing without adding XML declaration
    dom = xml.dom.minidom.parseString(xml_str)
    pretty_xml_str = dom.toprettyxml(indent='  ')
    
    # Remove the first line containing the unwanted XML declaration
    pretty_xml_str = '\n'.join(pretty_xml_str.split('\n')[1:])
    
    # Insert DOCTYPE declaration and XML declaration manually
    xml_declaration = '<?xml version="1.0" encoding="UTF-8"?>\n'
    doctype = '<!DOCTYPE task PUBLIC "-//OASIS//DTD DITA Task//EN" "task.dtd">\n'
    final_xml_str = xml_declaration + doctype + pretty_xml_str
    
    # Write formatted XML to the .dita file
    with open(dita_path, 'w', encoding='utf-8') as f:
        f.write(final_xml_str)

def main():
    if len(sys.argv) != 4:
        print("Usage: python script.py <input_docx_file> <output_dita_file> <task_id>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    dita_path = sys.argv[2]
    task_id = sys.argv[3]
    
    if not docx_path.lower().endswith('.docx'):
        print("Input file must be a .docx file.")
        sys.exit(1)
    
    if not dita_path.lower().endswith('.dita'):
        dita_path += '.dita'
    
    docx_to_dita_task(docx_path, dita_path, task_id)
    print(f"Conversion completed successfully. Output saved to {dita_path}")

if __name__ == '__main__':
    main()
