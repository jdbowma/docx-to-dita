import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import xml.etree.ElementTree as ET
import xml.dom.minidom
import os

def docx_to_dita_task(docx_path, dita_path, task_id, images_dir=None, ask_for_each_image=False):
    doc = Document(docx_path)
    
    # Create the root element of the DITA task document with task_id attribute
    root = ET.Element('task', id=task_id)
    
    # Add the title to the DITA document
    title = ET.SubElement(root, 'title')
    title.text = doc.paragraphs[0].text  # Assuming the first paragraph is the title
    
    # Add task body content
    task_body = ET.SubElement(root, 'taskbody')
    
    # Process paragraphs to identify steps, substeps, and images
    steps = ET.SubElement(task_body, 'steps')
    
    current_step = None
    current_substeps = None
    image_counter = 1
    
    for para in doc.paragraphs[1:]:  # Skipping the first paragraph (title)
        if para.style.name == 'List Paragraph':
            # New main step
            current_step = ET.SubElement(steps, 'step')
            step_cmd = ET.SubElement(current_step, 'cmd')
            step_cmd.text = para.text.strip()
            current_substeps = None
        elif para.style.name == 'List Number 2' and current_step is not None:
            # New substep
            if current_substeps is None:
                current_substeps = ET.SubElement(current_step, 'substeps')
            substep = ET.SubElement(current_substeps, 'substep')
            substep_cmd = ET.SubElement(substep, 'cmd')
            substep_cmd.text = para.text.strip()
        else:
            # Regular paragraphs are added as step info
            if current_step is not None:
                step_info = ET.SubElement(current_step, 'info')
                step_info.text = para.text.strip()
    
    # Handle images
    if images_dir or ask_for_each_image:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image = rel.target_part
                image_data = image.blob
                image_ext = os.path.splitext(image.partname)[-1]
                
                if ask_for_each_image:
                    image_filename = filedialog.asksaveasfilename(
                        defaultextension=image_ext,
                        filetypes=[("Image files", f"*{image_ext}")],
                        title=f"Save image {image_counter}"
                    )
                    if not image_filename:
                        continue
                else:
                    image_filename = f"image{image_counter}{image_ext}"
                    image_path = os.path.join(images_dir, image_filename)
                
                with open(image_path, 'wb') as img_file:
                    img_file.write(image_data)
                
                image_counter += 1
                
                if current_step is not None:
                    fig = ET.SubElement(current_step, 'fig')
                    title = ET.SubElement(fig, 'title')
                    title.text = ""  # Default to empty
                    image_ref = ET.SubElement(fig, 'image', href=image_filename)

    # Create a new XML tree
    tree = ET.ElementTree(root)
    
    # Convert XML to string without XML declaration
    xml_str = ET.tostring(root, encoding='utf-8', method='xml')
    
    # Parse XML string for pretty printing without adding XML declaration
    dom = xml.dom.minidom.parseString(xml_str)
    pretty_xml_str = dom.toprettyxml(indent='  ')
    
    # Remove the first line containing the unwanted XML declaration
    pretty_xml_str = '\n'.join(pretty_xml_str.split('\n')[1:])
    
    # Insert DOCTYPE declaration and XML declaration manually
    xml_declaration = '<?xml version="1.0" encoding="UTF-8"?>\n'
    doctype = '<!DOCTYPE task PUBLIC "-//OASIS//DTD DITA Task//EN" "task.dtd">\n'
    final_xml_str = xml_declaration + doctype + pretty_xml_str
    
    # Write formatted XML to the .dita file
    with open(dita_path, 'w', encoding='utf-8') as f:
        f.write(final_xml_str)

def browse_file():
    filename = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if filename:
        input_path_entry.delete(0, tk.END)
        input_path_entry.insert(0, filename)

def select_images_dir():
    directory = filedialog.askdirectory()
    if directory:
        images_dir_entry.delete(0, tk.END)
        images_dir_entry.insert(0, directory)

def convert_file():
    input_path = input_path_entry.get()
    output_path = output_path_entry.get()
    task_id = task_id_entry.get()
    images_dir = images_dir_entry.get() if images_checkbox_var.get() and not ask_each_image_var.get() else None
    ask_for_each_image = ask_each_image_var.get()
    
    if not input_path:
        messagebox.showerror("Error", "Please select an input file.")
        return
    
    if not output_path:
        messagebox.showerror("Error", "Please specify an output file.")
        return
    
    if not task_id:
        messagebox.showerror("Error", "Please enter a task ID.")
        return
    
    if not input_path.lower().endswith('.docx'):
        messagebox.showerror("Error", "Input file must be a .docx file.")
        return
    
    if not output_path.lower().endswith('.dita'):
        output_path += '.dita'
    
    try:
        if images_dir and not os.path.exists(images_dir):
            os.makedirs(images_dir)
        docx_to_dita_task(input_path, output_path, task_id, images_dir, ask_for_each_image)
        messagebox.showinfo("Success", f"Conversion completed successfully. Output saved to {output_path}")
        input_path_entry.delete(0, tk.END)
        output_path_entry.delete(0, tk.END)
        task_id_entry.delete(0, tk.END)
        images_dir_entry.delete(0, tk.END)
        images_checkbox_var.set(0)
        ask_each_image_var.set(0)
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred during conversion:\n{str(e)}")

def toggle_images_options():
    if images_checkbox_var.get():
        images_dir_label.grid(row=4, column=0, padx=10, pady=5, sticky=tk.W)
        images_dir_entry.grid(row=4, column=1, padx=10, pady=5)
        browse_images_dir_button.grid(row=4, column=2, padx=10, pady=5)
        ask_each_image_checkbox.grid(row=5, column=1, padx=10, pady=5, sticky=tk.W)
    else:
        images_dir_label.grid_remove()
        images_dir_entry.grid_remove()
        browse_images_dir_button.grid_remove()
        ask_each_image_checkbox.grid_remove()

# GUI setup
root = tk.Tk()
root.title("DOCX to DITA Converter")

# Input File
input_path_label = tk.Label(root, text="Input .docx file:")
input_path_label.grid(row=0, column=0, padx=10, pady=5, sticky=tk.W)

input_path_entry = tk.Entry(root, width=50)
input_path_entry.grid(row=0, column=1, padx=10, pady=5)

browse_button = tk.Button(root, text="Browse", command=browse_file)
browse_button.grid(row=0, column=2, padx=10, pady=5)

# Output File
output_path_label = tk.Label(root, text="Output .dita file:")
output_path_label.grid(row=1, column=0, padx=10, pady=5, sticky=tk.W)

output_path_entry = tk.Entry(root, width=50)
output_path_entry.grid(row=1, column=1, padx=10, pady=5)

# Task ID
task_id_label = tk.Label(root, text="Task ID:")
task_id_label.grid(row=2, column=0, padx=10, pady=5, sticky=tk.W)

task_id_entry = tk.Entry(root, width=50)
task_id_entry.grid(row=2, column=1, padx=10, pady=5)

# Images Checkbox
images_checkbox_var = tk.IntVar()
images_checkbox = tk.Checkbutton(root, text="Include images", variable=images_checkbox_var, command=toggle_images_options)
images_checkbox.grid(row=3, column=0, columnspan=2, padx=10, pady=5)

# Images Directory
images_dir_label = tk.Label(root, text="Images directory:")
images_dir_entry = tk.Entry(root, width=50)
browse_images_dir_button = tk.Button(root, text="Browse", command=select_images_dir)

# Ask for each image
ask_each_image_var = tk.IntVar()
ask_each_image_checkbox = tk.Checkbutton(root, text="Ask for each image", variable=ask_each_image_var)

# Convert Button
convert_button = tk.Button(root, text="Convert", command=convert_file)
convert_button.grid(row=6, column=1, padx=10, pady=10)

root.mainloop()
